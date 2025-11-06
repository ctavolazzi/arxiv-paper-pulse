# arxiv_paper_pulse/article_generator.py
"""
Simple article generator that fetches arXiv papers, analyzes them,
generates images, writes articles, and outputs DOCX files.
"""
from pathlib import Path
from datetime import datetime
import arxiv
from google import genai

from .documents import DocumentProcessor, DocumentInput, DocumentFromURL, DocumentProcessingConfig, OutputFormat
from .image_generator import ImageGenerator
from . import config


def _extract_paper_id(url_or_id: str) -> str:
    """Extract arXiv paper ID from URL or ID string."""
    if 'arxiv.org' in url_or_id:
        paper_id = url_or_id.split('/')[-1].replace('.pdf', '').replace('.abs', '')
    else:
        paper_id = url_or_id
    return paper_id


def _fetch_paper_metadata(paper_id: str):
    """Fetch paper metadata from arXiv."""
    search = arxiv.Search(id_list=[paper_id])
    client = arxiv.Client()
    results = list(client.results(search))
    if not results:
        raise ValueError(f"Paper {paper_id} not found on arXiv")
    paper = results[0]
    return {
        'title': paper.title,
        'authors': [author.name for author in paper.authors],
        'published': str(paper.published),
        'paper_id': paper_id,
        'arxiv_url': paper.entry_id
    }


def generate_article(url_or_id: str, output_format="docx") -> str:
    """
    Generate article from arXiv paper.

    Args:
        url_or_id: arXiv URL or paper ID (e.g., "2301.12345" or "https://arxiv.org/abs/2301.12345")
        output_format: "docx" (default) or "md"

    Returns:
        Path to generated article file
    """
    # Extract paper ID and fetch metadata
    paper_id = _extract_paper_id(url_or_id)
    metadata = _fetch_paper_metadata(paper_id)

    # Construct PDF URL (use format without .pdf extension to avoid redirects)
    pdf_url = f"https://arxiv.org/pdf/{paper_id}"

    # Initialize processors
    doc_processor = DocumentProcessor()
    img_generator = ImageGenerator()
    client = genai.Client(api_key=config.GEMINI_API_KEY)

    # Analyze PDF
    doc_input = DocumentInput(source=DocumentFromURL(url=pdf_url))
    doc_config = DocumentProcessingConfig(
        prompt="""Analyze this research paper comprehensively. Provide a detailed analysis covering:
1. Problem statement and significance
2. Methodology and approach
3. Key findings and contributions
4. Implications and applications
5. Limitations and future directions

Focus on clarity and depth. Extract key concepts that could be visualized.""",
        output_format=OutputFormat.TEXT
    )
    analysis_result = doc_processor.process(doc_input, doc_config)
    if not analysis_result.success:
        raise ValueError(f"Document analysis failed: {analysis_result.error}")
    analysis_text = analysis_result.text

    # Generate image prompt from analysis
    image_prompt_response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=[f"""Based on this paper analysis, create a detailed image prompt for a visual representation of the key concepts:

{analysis_text}

Generate a clear, descriptive image prompt that captures the essence of this research."""]
    )
    image_prompt = image_prompt_response.text.strip()

    # Generate and save image
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    image_filename = f"article_image_{paper_id}_{timestamp}.png"
    image_path = img_generator.generate_and_save(
        image_prompt,
        str(Path(config.IMAGE_OUTPUT_DIR) / image_filename)
    )

    # Generate article text
    article_prompt = f"""Write a comprehensive article about this research paper.

Paper Title: {metadata['title']}
Authors: {', '.join(metadata['authors'])}
Published: {metadata['published']}
arXiv ID: {metadata['paper_id']}

Paper Analysis:
{analysis_text}

Write an article with the following structure:
1. Title (use the paper title)
2. Abstract - Brief summary of the paper's contributions
3. Introduction - Context and motivation for the research
4. Deep Analysis - Detailed breakdown of the problem, methodology, findings, and implications
5. Visual Summary - Describe what the generated image represents
6. Conclusion - Summary of key takeaways and significance
7. References - Link to original paper

Write in a clear, accessible style suitable for readers interested in research."""

    article_response = client.models.generate_content(
        model="gemini-2.5-pro",
        contents=[article_prompt]
    )
    article_text = article_response.text
    print(f"   ✅ Article text generated ({len(article_text)} characters)")
    
    # Create output directory
    output_dir = Path(config.ARTICLE_OUTPUT_DIR)
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # Generate output file
    print(f"📝 Step 6/6: Assembling {output_format.upper()} document...")
    if output_format == "docx":
        from docx import Document
        from docx.shared import Inches

        doc = Document()

        # Title
        doc.add_heading(metadata['title'], 0)

        # Metadata
        metadata_para = doc.add_paragraph()
        metadata_para.add_run(f"Authors: {', '.join(metadata['authors'])}\n").bold = True
        metadata_para.add_run(f"Published: {metadata['published']}\n")
        metadata_para.add_run(f"arXiv ID: {metadata['paper_id']}\n")
        metadata_para.add_run(f"Original Paper: {metadata['arxiv_url']}\n")

        doc.add_paragraph()  # Spacing

        # Article content (split into paragraphs)
        for paragraph in article_text.split('\n\n'):
            if paragraph.strip():
                if paragraph.strip().startswith('#'):
                    # Heading
                    heading_text = paragraph.strip().lstrip('#').strip()
                    if paragraph.strip().startswith('###'):
                        doc.add_heading(heading_text, level=3)
                    elif paragraph.strip().startswith('##'):
                        doc.add_heading(heading_text, level=2)
                    else:
                        doc.add_heading(heading_text, level=1)
                else:
                    doc.add_paragraph(paragraph.strip())

        # Add image if exists
        if Path(image_path).exists():
            doc.add_paragraph()  # Spacing
            doc.add_heading("Visual Summary", level=2)
            doc.add_picture(image_path, width=Inches(5))
            doc.add_paragraph(f"*Image generated from paper analysis on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*").italic = True

        # Save
        output_filename = f"article_{paper_id}_{timestamp}.docx"
        output_path = output_dir / output_filename
        doc.save(str(output_path))
        print(f"   ✅ DOCX saved: {output_filename}")
        
    elif output_format == "md":
        # Markdown output
        markdown_content = f"""# {metadata['title']}

**Authors:** {', '.join(metadata['authors'])}
**Published:** {metadata['published']}
**arXiv ID:** {metadata['paper_id']}
**Original Paper:** {metadata['arxiv_url']}

---

{article_text}

---

## Generated Image

![Paper Analysis Image]({image_path})

*Image generated from paper analysis on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*
"""
        output_filename = f"article_{paper_id}_{timestamp}.md"
        output_path = output_dir / output_filename
        output_path.write_text(markdown_content, encoding='utf-8')
        print(f"   ✅ Markdown saved: {output_filename}")
    
    else:
        raise ValueError(f"Unsupported output format: {output_format}. Use 'docx' or 'md'")
    
    print()
    print(f"✅ Article generation complete!")
    print(f"   Output: {output_path}")
    print()
    
    return str(output_path)


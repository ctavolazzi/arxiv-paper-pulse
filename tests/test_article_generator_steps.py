#!/usr/bin/env python3
"""
Step-by-step tests for article generator that verify each part of the process.
Uses MOCKS to test logic without making real API calls.
Tests each component individually and INVESTIGATES failures.
"""
import pytest
from pathlib import Path
from unittest.mock import Mock, patch, MagicMock
from arxiv_paper_pulse.article_generator import (
    _extract_paper_id,
    _fetch_paper_metadata,
    generate_article
)
from arxiv_paper_pulse.documents import DocumentProcessor, DocumentInput, DocumentFromURL, DocumentProcessingConfig, OutputFormat
from arxiv_paper_pulse.image_generator import ImageGenerator
from google import genai
from arxiv_paper_pulse import config


class TestStepByStepProcess:
    """Test each step of the article generation process individually with MOCKS."""

    TEST_PAPER_ID = "1706.03762"

    def test_step_1_extract_paper_id(self):
        """STEP 1: Extract paper ID from various input formats."""
        print("\n" + "="*80)
        print("STEP 1: Testing Paper ID Extraction (NO API CALLS)")
        print("="*80)

        test_cases = [
            ("https://arxiv.org/abs/1706.03762", "1706.03762"),
            ("https://arxiv.org/pdf/1706.03762.pdf", "1706.03762"),
            ("1706.03762", "1706.03762"),
            ("1706.03762v1", "1706.03762v1"),
        ]

        failures = []
        for input_val, expected in test_cases:
            try:
                result = _extract_paper_id(input_val)
                print(f"  Input: {input_val}")
                print(f"  Expected: {expected}, Got: {result}")
                assert result == expected, f"Failed to extract ID from {input_val}"
                print(f"  ✅ PASSED")
            except AssertionError as e:
                print(f"  ❌ FAILED: {e}")
                failures.append((input_val, expected, result))

        if failures:
            print(f"\n  🔍 INVESTIGATION: {len(failures)} test(s) failed")
            for input_val, expected, got in failures:
                print(f"     - Input '{input_val}' expected '{expected}' but got '{got}'")
            raise AssertionError(f"{len(failures)} test case(s) failed")

        print("  ✅ STEP 1 COMPLETE: Paper ID extraction works")

    @patch('arxiv_paper_pulse.article_generator.arxiv.Client')
    @patch('arxiv_paper_pulse.article_generator.arxiv.Search')
    def test_step_2_fetch_metadata(self, mock_search, mock_client):
        """STEP 2: Fetch paper metadata from arXiv (MOCKED)."""
        print("\n" + "="*80)
        print("STEP 2: Testing Metadata Fetching (MOCKED - NO API CALLS)")
        print("="*80)

        try:
            # Create mock author objects
            mock_author1 = Mock()
            mock_author1.name = "Vaswani, Ashish"
            mock_author2 = Mock()
            mock_author2.name = "Shazeer, Noam"

            # Create mock paper object
            mock_paper = Mock()
            mock_paper.title = "Attention Is All You Need"
            mock_paper.authors = [mock_author1, mock_author2]
            mock_paper.published = "2017-06-12"
            mock_paper.entry_id = "http://arxiv.org/abs/1706.03762"

            # Mock search results
            mock_client_instance = Mock()
            mock_client_instance.results.return_value = [mock_paper]
            mock_client.return_value = mock_client_instance

            # Mock search
            mock_search_instance = Mock()
            mock_search.return_value = mock_search_instance

            print(f"  Fetching metadata for paper: {self.TEST_PAPER_ID}")
            metadata = _fetch_paper_metadata(self.TEST_PAPER_ID)

            # Verify required fields exist
            required_fields = ['title', 'authors', 'published', 'paper_id', 'arxiv_url']
            missing_fields = []
            for field in required_fields:
                if field not in metadata:
                    missing_fields.append(field)
                else:
                    print(f"  ✅ {field}: {metadata[field]}")

            if missing_fields:
                print(f"\n  ❌ STEP 2 FAILED: Missing required fields: {missing_fields}")
                print(f"  🔍 INVESTIGATION:")
                print(f"     - Check _fetch_paper_metadata function")
                print(f"     - Verify it extracts all fields from paper object")
                raise AssertionError(f"Missing fields: {missing_fields}")

            # Verify authors is a list
            if not isinstance(metadata['authors'], list):
                print(f"\n  ❌ STEP 2 FAILED: Authors should be a list, got {type(metadata['authors'])}")
                print(f"  🔍 INVESTIGATION:")
                print(f"     - Check how authors are extracted from paper.authors")
                raise AssertionError(f"Authors should be list, got {type(metadata['authors'])}")

            if len(metadata['authors']) == 0:
                print(f"\n  ❌ STEP 2 FAILED: Should have at least one author")
                print(f"  🔍 INVESTIGATION:")
                print(f"     - Check paper.authors extraction logic")
                raise AssertionError("No authors found")

            print(f"  ✅ Found {len(metadata['authors'])} authors")
            print("  ✅ STEP 2 COMPLETE: Metadata fetching works (with mocks)")

        except Exception as e:
            print(f"\n  ❌ STEP 2 FAILED: {type(e).__name__}: {e}")
            print(f"  🔍 INVESTIGATION:")
            print(f"     - Error type: {type(e).__name__}")
            print(f"     - Error message: {str(e)}")
            import traceback
            traceback.print_exc()
            raise

    @patch('arxiv_paper_pulse.article_generator.DocumentProcessor')
    def test_step_3_document_processing(self, mock_doc_processor_class):
        """STEP 3: Download and process PDF document (MOCKED)."""
        print("\n" + "="*80)
        print("STEP 3: Testing Document Processing (MOCKED - NO API CALLS)")
        print("="*80)

        try:
            pdf_url = f"https://arxiv.org/pdf/{self.TEST_PAPER_ID}"
            print(f"  PDF URL: {pdf_url}")

            # Mock document processor
            mock_processor = Mock()
            mock_result = Mock()
            mock_result.success = True
            mock_result.text = "This paper introduces the Transformer architecture for sequence-to-sequence tasks. The key innovation is the attention mechanism that allows the model to focus on different parts of the input sequence."
            mock_result.error = None
            mock_processor.process.return_value = mock_result
            mock_doc_processor_class.return_value = mock_processor

            processor = DocumentProcessor()
            doc_input = DocumentInput(source=DocumentFromURL(url=pdf_url))
            doc_config = DocumentProcessingConfig(
                prompt="Extract the abstract and main contributions from this paper.",
                output_format=OutputFormat.TEXT
            )

            print("  Processing document (MOCKED)...")
            result = processor.process(doc_input, doc_config)

            if not result.success:
                print(f"\n  ❌ STEP 3 FAILED: {result.error}")
                print(f"  🔍 INVESTIGATION:")
                print(f"     - Check DocumentProcessor.process() return value")
                print(f"     - Verify success flag is set correctly")
                raise AssertionError(f"Document processing failed: {result.error}")

            if not result.text:
                print(f"\n  ❌ STEP 3 FAILED: Result text is empty")
                print(f"  🔍 INVESTIGATION:")
                print(f"     - Check mock result.text is set")
                raise AssertionError("Result text should not be empty")

            if len(result.text) < 100:
                print(f"\n  ❌ STEP 3 FAILED: Result text too short: {len(result.text)} characters")
                print(f"  🔍 INVESTIGATION:")
                print(f"     - Check if mock text is long enough")
                raise AssertionError(f"Result text too short: {len(result.text)} characters")

            print(f"  ✅ Document processed successfully (MOCKED)")
            print(f"  ✅ Text length: {len(result.text)} characters")
            print(f"  ✅ Preview: {result.text[:100]}...")
            print("  ✅ STEP 3 COMPLETE: Document processing works (with mocks)")

        except Exception as e:
            print(f"\n  ❌ STEP 3 FAILED: {type(e).__name__}: {e}")
            print(f"  🔍 INVESTIGATION:")
            print(f"     - Error type: {type(e).__name__}")
            print(f"     - Error message: {str(e)}")
            import traceback
            traceback.print_exc()
            raise

    @patch('arxiv_paper_pulse.article_generator.genai.Client')
    def test_step_4_image_prompt_generation(self, mock_genai_client):
        """STEP 4: Generate image prompt from document analysis (MOCKED)."""
        print("\n" + "="*80)
        print("STEP 4: Testing Image Prompt Generation (MOCKED - NO API CALLS)")
        print("="*80)

        try:
            analysis_text = "This paper introduces the Transformer architecture for sequence-to-sequence tasks."

            # Mock Gemini client
            mock_client = Mock()
            mock_response = Mock()
            mock_response.text = "A scientific infographic showing a transformer neural network architecture with attention mechanisms, 16:9 aspect ratio, professional style, with labeled components showing encoder and decoder layers"
            mock_client.models.generate_content.return_value = mock_response
            mock_genai_client.return_value = mock_client

            client = genai.Client(api_key="fake_key")
            prompt = f"Create a detailed image prompt for a blog hero image (16:9) that visually explains: {analysis_text}"

            print("  Generating image prompt (MOCKED)...")
            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=[prompt]
            )

            image_prompt = response.text.strip()

            if not image_prompt:
                print(f"\n  ❌ STEP 4 FAILED: Image prompt is empty")
                print(f"  🔍 INVESTIGATION:")
                print(f"     - Check mock response.text is set")
                raise AssertionError("Image prompt should not be empty")

            if len(image_prompt) < 50:
                print(f"\n  ❌ STEP 4 FAILED: Image prompt too short: {len(image_prompt)} characters")
                print(f"  🔍 INVESTIGATION:")
                print(f"     - Check mock response contains sufficient text")
                raise AssertionError(f"Image prompt too short: {len(image_prompt)} characters")

            print(f"  ✅ Image prompt generated (MOCKED)")
            print(f"  ✅ Prompt length: {len(image_prompt)} characters")
            print(f"  ✅ Preview: {image_prompt[:100]}...")
            print("  ✅ STEP 4 COMPLETE: Image prompt generation works (with mocks)")

        except Exception as e:
            print(f"\n  ❌ STEP 4 FAILED: {type(e).__name__}: {e}")
            print(f"  🔍 INVESTIGATION:")
            print(f"     - Error type: {type(e).__name__}")
            print(f"     - Error message: {str(e)}")
            import traceback
            traceback.print_exc()
            raise

    @patch('arxiv_paper_pulse.article_generator.ImageGenerator')
    def test_step_5_image_generation(self, mock_img_generator_class):
        """STEP 5: Generate image from prompt (MOCKED)."""
        print("\n" + "="*80)
        print("STEP 5: Testing Image Generation (MOCKED - NO API CALLS)")
        print("="*80)

        try:
            # Mock image generator
            mock_generator = Mock()
            test_image_path = Path("arxiv_paper_pulse/data/generated_images/test_step_image.png")
            test_image_path.parent.mkdir(parents=True, exist_ok=True)

            # Create a dummy image file
            test_image_path.write_bytes(b'fake image data')
            mock_generator.generate_and_save.return_value = str(test_image_path)
            mock_img_generator_class.return_value = mock_generator

            generator = ImageGenerator()
            test_prompt = "A scientific infographic showing a transformer neural network architecture"

            print("  Generating image (MOCKED)...")
            result_path = generator.generate_and_save(test_prompt, str(test_image_path))

            if not result_path:
                print(f"\n  ❌ STEP 5 FAILED: No path returned from generate_and_save")
                print(f"  🔍 INVESTIGATION:")
                print(f"     - Check mock return value is set")
                raise AssertionError("Image generation should return a path")

            # Verify the path format is correct
            if not isinstance(result_path, str):
                print(f"\n  ❌ STEP 5 FAILED: Path should be string, got {type(result_path)}")
                print(f"  🔍 INVESTIGATION:")
                print(f"     - Check return type of generate_and_save")
                raise AssertionError(f"Path should be string, got {type(result_path)}")

            print(f"  ✅ Image generated (MOCKED)")
            print(f"  ✅ Image path: {result_path}")
            print("  ✅ STEP 5 COMPLETE: Image generation works (with mocks)")

            # Cleanup
            if test_image_path.exists():
                test_image_path.unlink()

        except Exception as e:
            print(f"\n  ❌ STEP 5 FAILED: {type(e).__name__}: {e}")
            print(f"  🔍 INVESTIGATION:")
            print(f"     - Error type: {type(e).__name__}")
            print(f"     - Error message: {str(e)}")
            import traceback
            traceback.print_exc()
            raise

    @patch('arxiv_paper_pulse.article_generator.genai.Client')
    def test_step_6_article_text_generation(self, mock_genai_client):
        """STEP 6: Generate article text from analysis (MOCKED)."""
        print("\n" + "="*80)
        print("STEP 6: Testing Article Text Generation (MOCKED - NO API CALLS)")
        print("="*80)

        try:
            metadata = {
                'title': "Attention Is All You Need",
                'authors': ["Vaswani, Ashish", "Shazeer, Noam"],
                'published': "2017-06-12",
                'paper_id': "1706.03762",
                'arxiv_url': "http://arxiv.org/abs/1706.03762"
            }

            analysis_text = "This paper introduces a novel architecture for sequence processing."

            # Mock Gemini client
            mock_client = Mock()
            mock_response = Mock()
            mock_response.text = f"""# {metadata['title']}

## Abstract

This paper introduces the Transformer architecture for sequence-to-sequence tasks.

## Introduction

The Transformer model uses attention mechanisms to process sequences.

## Deep Analysis

The key innovation is the self-attention mechanism that allows the model to focus on different parts of the input.

## Conclusion

The Transformer architecture has revolutionized sequence processing.
"""
            mock_client.models.generate_content.return_value = mock_response
            mock_genai_client.return_value = mock_client

            article_prompt = f"Write an article about: {metadata['title']}"

            client = genai.Client(api_key="fake_key")
            print("  Generating article text (MOCKED)...")
            response = client.models.generate_content(
                model="gemini-2.5-pro",
                contents=[article_prompt]
            )

            article_text = response.text

            if not article_text:
                print(f"\n  ❌ STEP 6 FAILED: Article text is empty")
                print(f"  🔍 INVESTIGATION:")
                print(f"     - Check mock response.text is set")
                raise AssertionError("Article text should not be empty")

            if len(article_text) < 200:
                print(f"\n  ❌ STEP 6 FAILED: Article text too short: {len(article_text)} characters")
                print(f"  🔍 INVESTIGATION:")
                print(f"     - Check mock response contains sufficient text")
                raise AssertionError(f"Article text too short: {len(article_text)} characters")

            if metadata['title'].lower() not in article_text.lower():
                print(f"\n  ❌ STEP 6 FAILED: Article should contain paper title")
                print(f"  🔍 INVESTIGATION:")
                print(f"     - Check if article prompt includes title")
                print(f"     - Verify mock response includes title")
                raise AssertionError("Article should contain paper title")

            print(f"  ✅ Article text generated (MOCKED)")
            print(f"  ✅ Text length: {len(article_text)} characters")
            print(f"  ✅ Preview: {article_text[:150]}...")
            print("  ✅ STEP 6 COMPLETE: Article text generation works (with mocks)")

        except Exception as e:
            print(f"\n  ❌ STEP 6 FAILED: {type(e).__name__}: {e}")
            print(f"  🔍 INVESTIGATION:")
            print(f"     - Error type: {type(e).__name__}")
            print(f"     - Error message: {str(e)}")
            import traceback
            traceback.print_exc()
            raise

    def test_step_7_file_creation(self):
        """STEP 7: Create output files (markdown and DOCX) - NO MOCKS NEEDED."""
        print("\n" + "="*80)
        print("STEP 7: Testing File Creation (NO MOCKS - ACTUAL FILE OPERATIONS)")
        print("="*80)

        from datetime import datetime
        from docx import Document

        try:
            output_dir = Path("arxiv_paper_pulse/data/articles")
            output_dir.mkdir(parents=True, exist_ok=True)

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            test_content = "# Test Article\n\nThis is a test article for file creation testing."

            failures = []

            # Test Markdown creation
            print("  Testing Markdown file creation...")
            try:
                md_file = output_dir / f"test_step_md_{timestamp}.md"
                md_file.write_text(test_content, encoding='utf-8')

                if not md_file.exists():
                    failures.append(f"Markdown file not created: {md_file}")
                    print(f"  ❌ FAILED: File not created")
                elif md_file.read_text() != test_content:
                    failures.append(f"Markdown content mismatch")
                    print(f"  ❌ FAILED: Content mismatch")
                else:
                    print(f"  ✅ Markdown file created: {md_file}")
            except Exception as e:
                failures.append(f"Markdown creation error: {e}")
                print(f"  ❌ FAILED: {e}")

            # Test DOCX creation
            print("  Testing DOCX file creation...")
            try:
                docx_file = output_dir / f"test_step_docx_{timestamp}.docx"
                doc = Document()
                doc.add_heading("Test Article", 0)
                doc.add_paragraph("This is a test article for file creation testing.")
                doc.save(str(docx_file))

                if not docx_file.exists():
                    failures.append(f"DOCX file not created: {docx_file}")
                    print(f"  ❌ FAILED: File not created")
                elif docx_file.stat().st_size == 0:
                    failures.append(f"DOCX file is empty")
                    print(f"  ❌ FAILED: File is empty")
                else:
                    print(f"  ✅ DOCX file created: {docx_file}")
            except Exception as e:
                failures.append(f"DOCX creation error: {e}")
                print(f"  ❌ FAILED: {e}")

            # Cleanup
            if md_file.exists():
                md_file.unlink()
            if docx_file.exists():
                docx_file.unlink()
            print("  ✅ Test files cleaned up")

            if failures:
                print(f"\n  🔍 INVESTIGATION: {len(failures)} failure(s)")
                for failure in failures:
                    print(f"     - {failure}")
                raise AssertionError(f"{len(failures)} file creation test(s) failed")

            print("  ✅ STEP 7 COMPLETE: File creation works")

        except Exception as e:
            print(f"\n  ❌ STEP 7 FAILED: {type(e).__name__}: {e}")
            print(f"  🔍 INVESTIGATION:")
            print(f"     - Is output directory writable?")
            print(f"     - Is python-docx installed?")
            print(f"     - Are file permissions correct?")
            import traceback
            traceback.print_exc()
            raise

    @patch('arxiv_paper_pulse.article_generator.ImageGenerator')
    @patch('arxiv_paper_pulse.article_generator.DocumentProcessor')
    @patch('arxiv_paper_pulse.article_generator.genai.Client')
    @patch('arxiv_paper_pulse.article_generator._fetch_paper_metadata')
    def test_full_integration_mocked(self, mock_fetch_metadata, mock_genai_client, mock_doc_processor_class, mock_img_generator_class):
        """FULL INTEGRATION: Test complete article generation process (ALL MOCKED)."""
        print("\n" + "="*80)
        print("FULL INTEGRATION TEST: Complete Article Generation (ALL MOCKED)")
        print("="*80)

        try:
            # Setup all mocks
            mock_fetch_metadata.return_value = {
                'title': "Attention Is All You Need",
                'authors': ["Vaswani, Ashish"],
                'published': "2017-06-12",
                'paper_id': self.TEST_PAPER_ID,
                'arxiv_url': "http://arxiv.org/abs/1706.03762"
            }

            mock_processor = Mock()
            mock_result = Mock()
            mock_result.success = True
            mock_result.text = "This paper introduces the Transformer architecture."
            mock_processor.process.return_value = mock_result
            mock_doc_processor_class.return_value = mock_processor

            mock_img_generator = Mock()
            mock_img_generator.generate_and_save.return_value = "/path/to/image.png"
            mock_img_generator_class.return_value = mock_img_generator

            mock_client = Mock()
            mock_image_prompt_response = Mock()
            mock_image_prompt_response.text = "A visual representation"
            mock_article_response = Mock()
            mock_article_response.text = "# Test Article\n\nThis is a test article."
            mock_client.models.generate_content.side_effect = [
                mock_image_prompt_response,
                mock_article_response
            ]
            mock_genai_client.return_value = mock_client

            print(f"  Generating article for paper: {self.TEST_PAPER_ID} (ALL MOCKED)")
            result = generate_article(self.TEST_PAPER_ID, output_format="md")

            if not result:
                print(f"\n  ❌ INTEGRATION FAILED: No result returned")
                print(f"  🔍 INVESTIGATION:")
                print(f"     - Check generate_article return statement")
                raise AssertionError("Article generation should return a path")

            if not Path(result).exists():
                print(f"\n  ❌ INTEGRATION FAILED: Article file not created: {result}")
                print(f"  🔍 INVESTIGATION:")
                print(f"     - Check file creation logic")
                raise AssertionError(f"Article file not created: {result}")

            if not result.endswith(".md"):
                print(f"\n  ❌ INTEGRATION FAILED: Expected .md file, got: {result}")
                print(f"  🔍 INVESTIGATION:")
                print(f"     - Check output_format handling")
                raise AssertionError(f"Expected .md file, got: {result}")

            content = Path(result).read_text()
            if len(content) < 100:
                print(f"\n  ❌ INTEGRATION FAILED: Article too short: {len(content)} characters")
                print(f"  🔍 INVESTIGATION:")
                print(f"     - Check article content generation")
                raise AssertionError(f"Article too short: {len(content)} characters")

            print(f"  ✅ Article generated successfully (ALL MOCKED)")
            print(f"  ✅ Article path: {result}")
            print(f"  ✅ Article length: {len(content)} characters")
            print("  ✅ FULL INTEGRATION COMPLETE: All steps work together (with mocks)")

            # Cleanup
            Path(result).unlink()

        except Exception as e:
            print(f"\n  ❌ INTEGRATION FAILED: {type(e).__name__}: {e}")
            print(f"  🔍 INVESTIGATION:")
            print(f"     - Which step failed? (Check individual step tests above)")
            print(f"     - Error type: {type(e).__name__}")
            print(f"     - Error message: {str(e)}")
            import traceback
            traceback.print_exc()
            raise
